from tkinter.messagebox import showinfo, askyesno

import win32com
from docx import Document
from tkinter import *
import win32print
import win32api
from win32com import client
import customtkinter


def print_bmi():
    try:
        word = win32com.client.Dispatch("Word.Application")

        doc = word.Documents.Open(str("//192.168.0.220/Shared PTO1/Уляшкин К.А/Гаврики"
                                      "/документы по жильцам/Квартал 26/1Оформление ТУ/ТУ Фулл Пакет.docx"))
        doc.SaveAs(str("//192.168.0.220/Shared PTO1/Уляшкин К.А/Гаврики"
                       "/документы по жильцам/Квартал 26/1Оформление ТУ/ТУ Фулл Пакет.pdf"), FileFormat=17)

        # doc = word.Documents.Open(str("C:/Users/User/PycharmProjects/pythonProject1/example.docx"))
        # doc.SaveAs(str("C:/Users/User/PycharmProjects/pythonProject1/example.pdf"), FileFormat=17)
        # doc.Close(0) "//192.168.0.220/Shared PTO1/Уляшкин К.А/Гаврики/документы
        # по жильцам/Квартал 30 ТХ/1Оформление ТУ"

        nameprinter = win32print.GetDefaultPrinter()
        defaultprinter = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}
        handle = win32print.OpenPrinter(nameprinter, defaultprinter)

        # ## Если изменить level на другое число, то не сработает
        level = 2
        # ## Получаем значения принтера
        attributes = win32print.GetPrinter(handle, level)
        # ## Настройка двухсторонней печати
        attributes['pDevMode'].Duplex = 2  # flip over  3 - это короткий 2 - это длинный край
        # ## Передаем нужные значения в принтер
        # # win32print.SetPrinter(handle, level, attributes, 0)
        # # win32print.GetPrinter(handle, level)['pDevMode'].Duplex
        # ## Предупреждаем принтер о старте печати
        # win32print.StartDocPrinter(handle, 1, ["example.pdf", None, "raw"])
        # 2 в начале для открытия pdf и его сворачивания, для открытия без сворачивания поменяйте на 1
        win32api.ShellExecute(2, 'print', "ТУ Фулл Пакет.pdf", '.', '/manualstoprint', 0)
        # "Закрываем" принтер
        win32print.ClosePrinter(handle)
    except Exception as err:
        my_file = open("logs.txt", "w+")
        my_file.write(str(err))
        my_file.close()


def clickPrint():
    result = askyesno(title="Подтвержение операции", message="Документ будет отпрвлен на печать.Продолжить?")
    if result:
        print_bmi()
    else:
        showinfo("Результат", "Операция отменена")


def clickSave():
    result = askyesno(title="Подтвержение операции", message="Подтвердить операцию?")
    if result:
        calculate_bmi()
    else:
        showinfo("Результат", "Операция отменена")


def calculate_bmi():
    if enabled.get() == 1:
        document = Document("380/ТУ ФуллПакет 380.docx")
    else:
        document = Document("220/ТУ ФуллПакет 220.docx")

    document.core_properties.author = name.get()  # Инициалы
    document.core_properties.category = currentDate.get()  # Дата
    document.core_properties.comments = passport.get()  # адрес
    document.core_properties.content_status = pyNumber.get()  # Номер счетчика
    document.core_properties.keywords = kadNumber.get()  # кадастровый номер
    document.core_properties.title = address.get()  # Имя и паспортные данные
    document.core_properties.subject = opNumber.get()  # Номер опоры

    cell = document.tables[6].rows[3].cells[2]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_text(modelPY.get())

    cell = document.tables[6].rows[11].cells[2]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_text(datePY.get())

    cell = document.tables[6].rows[12].cells[2]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_text("4 кв " + datePY.get())

    cell = document.tables[6].rows[20].cells[2]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_text("T - " + pokazaniya.get())

    document.paragraphs[
        147].text = "'8. Наложены пломбы:  На электросчетчике  №" + \
                    plomba1.get() + " На вводном автомате №" + plomba2.get()

    document.save("ТУ Фулл Пакет.docx")


customtkinter.set_appearance_mode("dark")
customtkinter.set_default_color_theme("dark-blue")

window = customtkinter.CTk()
window.title("Оформление технических условий")
window.geometry('650x600')

frame = customtkinter.CTkFrame(
    window,  # Обязательный параметр, который указывает окно для размещения Frame.
    width=650,
    height=450
)

frame.pack(pady=0, padx=0, expand=True)

# Адрес
address = customtkinter.CTkLabel(
    frame,
    text="Введите адрес  ",
    font=("Roboto", 20)
)
address.grid(row=1, column=1, padx=20, pady=(10, 10))

address = customtkinter.CTkEntry(
    frame
)
address.grid(row=1, column=2)


# Пасспорт
passport = customtkinter.CTkLabel(
    frame,
    text="Введите паспортные данные  ",
    font=("Roboto", 20)
)
passport.grid(row=2, column=1, padx=20, pady=(10, 10))

passport = customtkinter.CTkEntry(
    frame
)
passport.grid(row=2, column=2, pady=5)

# Инициалы
name = customtkinter.CTkLabel(
    frame,
    text="Инициалы",
    font=("Roboto", 20)
)
name.grid(row=3, column=1)

name = customtkinter.CTkEntry(
    frame
)
name.grid(row=3, column=2, pady=5)

# Дата
currentDate = customtkinter.CTkLabel(
    frame,
    text="Введите текущую дату",
    font=("Roboto", 20)
)
currentDate.grid(row=4, column=1)

currentDate = customtkinter.CTkEntry(
    frame
)
currentDate.grid(row=4, column=2, pady=5)

# Номер счетчика
pyNumber = customtkinter.CTkLabel(
    frame,
    text="Введите номер счетчика",
    font=("Roboto", 20)
)
pyNumber.grid(row=5, column=1)

pyNumber = customtkinter.CTkEntry(
    frame
)
pyNumber.grid(row=5, column=2, pady=5)

# Номер опоры
opNumber = customtkinter.CTkLabel(
    frame,
    text="Введите номер опоры",
    font=("Roboto", 20)
)
opNumber.grid(row=6, column=1)

opNumber = customtkinter.CTkEntry(
    frame
)
opNumber.grid(row=6, column=2, pady=5)

# Кадастровый номер
kadNumber = customtkinter.CTkLabel(
    frame,
    text="Введите кадастровый номер  ",
    font=("Roboto", 20)
)
kadNumber.grid(row=7, column=1)

kadNumber = customtkinter.CTkEntry(
    frame
)
kadNumber.grid(row=7, column=2, pady=5)

# Пломба1
plomba1 = customtkinter.CTkLabel(
    frame,
    text="Пломба 1",
    font=("Roboto", 20)
)
plomba1.grid(row=8, column=1)

plomba1 = customtkinter.CTkEntry(
    frame
)
plomba1.grid(row=8, column=2, pady=5)

# Пломба2
plomba2 = customtkinter.CTkLabel(
    frame,
    text="Пломба 2",
    font=("Roboto", 20)
)
plomba2.grid(row=9, column=1)

plomba2 = customtkinter.CTkEntry(
    frame
)
plomba2.grid(row=9, column=2, pady=5)

# modelPY
modelPY = customtkinter.CTkLabel(
    frame,
    text="Модель счетчика",
    font=("Roboto", 20)
)
modelPY.grid(row=10, column=1)

modelPY = customtkinter.CTkEntry(
    frame
)
modelPY.grid(row=10, column=2, pady=5)

# godVipuska
datePY = customtkinter.CTkLabel(
    frame,
    text="Год выпуска",
    font=("Roboto", 20)
)
datePY.grid(row=11, column=1)

datePY = customtkinter.CTkEntry(
    frame
)
datePY.grid(row=11, column=2, pady=5)

# pokazaniya
pokazaniya = customtkinter.CTkLabel(
    frame,
    text="Показания",
    font=("Roboto", 20)
)
pokazaniya.grid(row=12, column=1)

pokazaniya = customtkinter.CTkEntry(
    frame
)
pokazaniya.grid(row=12, column=2, pady=5)

# Buttons
enabled = IntVar()
example = customtkinter.CTkCheckBox(
    frame,
    text='380',
    variable=enabled
)
example.grid(row=13, column=1, padx=(10, 10))

cal_btn = customtkinter.CTkButton(
    frame,
    text='Cохранить в файл',
    command=clickSave
)
cal_btn.grid(row=15, column=1, padx=20, pady=(20, 15))

print_btn = customtkinter.CTkButton(
    frame,
    text='Печать',
    command=clickPrint
)
print_btn.grid(row=15, column=2, padx=20, pady=(20, 15))

window.mainloop()
